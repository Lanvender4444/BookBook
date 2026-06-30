"""RAG 服务：文档解析、分块、向量化（Embedding API）、混合检索（BM25 + 向量 + RRF 重排）。

设计要点：
- 向量用激活的 LLM 厂商 embedding 接口；厂商无 embedding（如 Anthropic）时只走 BM25。
- 检索后端分两条路，按可用性自动选择：
    A. LanceDB（嵌入式向量库，全平台）：原生 BM25(全文) + 向量 + 内置 RRF 混合重排；
    B. 纯 Python 回落：BM25 候选 + SQLite 内嵌向量暴力余弦候选 + 手写 RRF 融合。
  任一缺失都不影响功能。
- 四类知识源：writing_guide / style / reference / continuation。
- 续写(continuation)特殊处理：除检索外固定携带文档"尾部"内容，用于生成后续情节。
"""

import math
import re
import json
import struct
import uuid
from pathlib import Path

import httpx
from database import SessionLocal
from models import KnowledgeSource, KnowledgeChunk, ActiveModel, ProviderConfig
from config import PROVIDERS
from services import vector_store

CATEGORIES = ("writing_guide", "style", "reference", "continuation")

CHUNK_SIZE = 1000
CHUNK_OVERLAP = 150
EMBED_BATCH = 32
EMBED_TIMEOUT = httpx.Timeout(60.0, connect=15.0)

# 各厂商 embedding 模型映射；不在表内的厂商 → BM25 降级
EMBEDDING_MODELS = {
    "openai": "text-embedding-3-small",
    "zhipu": "embedding-3",
    "qwen": "text-embedding-v3",
    "qwen_intl": "text-embedding-v3",
    "gemini": "text-embedding-004",
    "siliconflow": "BAAI/bge-m3",
    "siliconflow_intl": "BAAI/bge-m3",
    "mistral": "mistral-embed",
    "nvidia": "nvidia/nv-embed-v1",
    "together": "BAAI/bge-large-en-v1.5",
}


# ---------------------------------------------------------------- 文本提取

class ExtractError(Exception):
    pass


def extract_text(file_path: str) -> str:
    path = Path(file_path)
    if not path.exists():
        raise ExtractError(f"文件不存在: {file_path}")
    suffix = path.suffix.lower()
    if suffix in (".txt", ".md", ".markdown", ".text"):
        return _read_text_file(path)
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix == ".epub":
        return _extract_epub(path)
    raise ExtractError(f"不支持的文件格式: {suffix}（支持 txt/md/pdf/docx/epub）")


def _read_text_file(path: Path) -> str:
    raw = path.read_bytes()
    for enc in ("utf-8", "utf-8-sig", "gb18030", "big5", "latin-1"):
        try:
            return raw.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue
    return raw.decode("utf-8", errors="replace")


def _extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise ExtractError("缺少 pypdf 依赖") from e
    reader = PdfReader(str(path))
    pages = [(page.extract_text() or "") for page in reader.pages]
    text = "\n\n".join(pages).strip()
    if not text:
        raise ExtractError("PDF 未提取到文本（可能是扫描件）")
    return text


def _extract_docx(path: Path) -> str:
    try:
        import docx
    except ImportError as e:
        raise ExtractError("缺少 python-docx 依赖") from e
    doc = docx.Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(p for p in parts if p and p.strip())


_TAG_RE = re.compile(r"<[^>]+>")


def _extract_epub(path: Path) -> str:
    try:
        from ebooklib import epub, ITEM_DOCUMENT
    except ImportError as e:
        raise ExtractError("缺少 ebooklib 依赖") from e
    book = epub.read_epub(str(path))
    parts = []
    for item in book.get_items_of_type(ITEM_DOCUMENT):
        html = item.get_content().decode("utf-8", errors="replace")
        html = re.sub(r"<(script|style)[^>]*>.*?</\1>", " ", html, flags=re.S | re.I)
        html = re.sub(r"</(p|div|h[1-6]|li|br)>", "\n", html, flags=re.I)
        text = _TAG_RE.sub(" ", html)
        text = re.sub(r"[ \t]+", " ", text)
        parts.append(text.strip())
    return "\n\n".join(p for p in parts if p)


def get_source_text(source: KnowledgeSource) -> str:
    """取知识源全文：纯文本投递直接返回；文件型走解析。"""
    if source.text_content:
        return source.text_content
    if source.file_path:
        return extract_text(source.file_path)
    raise ExtractError("知识源既无文本内容也无文件路径")


# ---------------------------------------------------------------- 分块

def chunk_text(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    text = re.sub(r"\r\n?", "\n", text).strip()
    if not text:
        return []
    paragraphs = [p.strip() for p in re.split(r"\n{2,}", text) if p.strip()]
    chunks, current = [], ""
    for para in paragraphs:
        # 超长段落内部硬切
        while len(para) > size:
            head, para = para[:size], para[max(0, size - overlap):]
            if current:
                chunks.append(current)
                current = ""
            chunks.append(head)
        if len(current) + len(para) + 1 <= size:
            current = f"{current}\n{para}".strip()
        else:
            if current:
                chunks.append(current)
            # 与上一块保留少量重叠
            tail = chunks[-1][-overlap:] if chunks else ""
            current = (tail + "\n" + para).strip() if tail else para
    if current:
        chunks.append(current)
    return chunks


# ---------------------------------------------------------------- Embedding

def _get_embedding_provider():
    """返回 (provider_id, api_key, base_url, embed_model)；不可用返回 None。"""
    db = SessionLocal()
    try:
        active = db.query(ActiveModel).order_by(ActiveModel.id.desc()).first()
        if not active:
            return None
        embed_model = EMBEDDING_MODELS.get(active.provider_id)
        if not embed_model:
            return None
        config = (
            db.query(ProviderConfig)
            .filter(ProviderConfig.provider_id == active.provider_id)
            .first()
        )
        if not config or not config.api_key:
            return None
        from services.llm_service import _decrypt_api_key

        api_key = _decrypt_api_key(config.api_key)
        base_url = config.base_url or PROVIDERS.get(active.provider_id, {}).get(
            "default_base_url", ""
        )
        if not api_key:
            return None
        return (active.provider_id, api_key, base_url, embed_model)
    finally:
        db.close()


def embed_texts(texts: list[str]) -> tuple[list[list[float]], str] | None:
    """批量向量化。成功返回 (vectors, model_tag)；无可用 embedding 服务返回 None。"""
    provider = _get_embedding_provider()
    if not provider:
        return None
    provider_id, api_key, base_url, model = provider
    vectors = []
    with httpx.Client(timeout=EMBED_TIMEOUT) as client:
        for i in range(0, len(texts), EMBED_BATCH):
            batch = texts[i : i + EMBED_BATCH]
            if provider_id == "gemini":
                vectors.extend(_embed_gemini(client, api_key, model, batch))
            else:
                vectors.extend(_embed_openai_compatible(client, api_key, base_url, model, batch))
    return vectors, f"{provider_id}:{model}"


def _embed_openai_compatible(client, api_key, base_url, model, batch):
    resp = client.post(
        f"{base_url.rstrip('/')}/embeddings",
        headers={"Authorization": f"Bearer {api_key}"},
        json={"model": model, "input": batch},
    )
    resp.raise_for_status()
    data = sorted(resp.json()["data"], key=lambda d: d["index"])
    return [d["embedding"] for d in data]


def _embed_gemini(client, api_key, model, batch):
    url = (
        "https://generativelanguage.googleapis.com/v1beta/"
        f"models/{model}:batchEmbedContents?key={api_key}"
    )
    body = {
        "requests": [
            {"model": f"models/{model}", "content": {"parts": [{"text": t}]}}
            for t in batch
        ]
    }
    resp = client.post(url, json=body)
    resp.raise_for_status()
    return [e["values"] for e in resp.json()["embeddings"]]


def _vec_to_bytes(vec: list[float]) -> bytes:
    return struct.pack(f"{len(vec)}f", *vec)


def _bytes_to_vec(blob: bytes) -> list[float]:
    return list(struct.unpack(f"{len(blob) // 4}f", blob))


def _cosine(a: list[float], b: list[float]) -> float:
    dot = sum(x * y for x, y in zip(a, b))
    na = math.sqrt(sum(x * x for x in a))
    nb = math.sqrt(sum(x * x for x in b))
    if na == 0 or nb == 0:
        return 0.0
    return dot / (na * nb)


# ---------------------------------------------------------------- BM25 降级

_WORD_RE = re.compile(r"[a-zA-Z0-9_]+")
_CJK_RE = re.compile(r"[一-鿿぀-ヿ가-힯]")


def _tokenize(text: str) -> list[str]:
    """轻量分词：拉丁词 + CJK 二元组，零依赖。"""
    tokens = [w.lower() for w in _WORD_RE.findall(text)]
    cjk = _CJK_RE.findall(text)
    tokens.extend("".join(pair) for pair in zip(cjk, cjk[1:]))
    tokens.extend(cjk)
    return tokens


def _bm25_rank(query: str, docs: list[str], k1: float = 1.5, b: float = 0.75) -> list[float]:
    tokenized = [_tokenize(d) for d in docs]
    n = len(docs)
    if n == 0:
        return []
    avgdl = sum(len(t) for t in tokenized) / n or 1.0
    df = {}
    for toks in tokenized:
        for t in set(toks):
            df[t] = df.get(t, 0) + 1
    q_tokens = _tokenize(query)
    scores = []
    for toks in tokenized:
        tf = {}
        for t in toks:
            tf[t] = tf.get(t, 0) + 1
        score = 0.0
        for qt in q_tokens:
            if qt not in tf:
                continue
            idf = math.log(1 + (n - df[qt] + 0.5) / (df[qt] + 0.5))
            denom = tf[qt] + k1 * (1 - b + b * len(toks) / avgdl)
            score += idf * tf[qt] * (k1 + 1) / denom
        scores.append(score)
    return scores


# ---------------------------------------------------------------- 索引

def index_source(source_id: str) -> dict:
    """解析 → 分块 → 向量化（可用时）→ 入库。返回状态摘要。"""
    db = SessionLocal()
    try:
        source = db.query(KnowledgeSource).filter(KnowledgeSource.id == source_id).first()
        if not source:
            raise ValueError(f"知识源不存在: {source_id}")
        source.index_status = "indexing"
        source.index_error = None
        db.commit()
        try:
            text = get_source_text(source)
            chunks = chunk_text(text)
            if not chunks:
                raise ExtractError("文档内容为空")

            # 清旧块（SQLite 文本/元数据）与旧向量/全文索引（LanceDB）
            db.query(KnowledgeChunk).filter(KnowledgeChunk.source_id == source_id).delete()
            vector_store.delete_by_source(source_id)

            embedded = None
            try:
                embedded = embed_texts(chunks)
            except Exception as e:
                # embedding 失败不致命，降级 BM25
                print(f"[RAG] embedding failed, fallback to BM25: {e}")

            # 先建 SQLite 块（文本 + 元数据，BM25 与命中回显都靠它），flush 拿自增主键
            rows = []
            for i, chunk in enumerate(chunks):
                row = KnowledgeChunk(source_id=source_id, chunk_index=i, text=chunk)
                db.add(row)
                rows.append(row)
            db.flush()

            status = "bm25"
            if embedded:
                vectors, model = embedded
                wrote_milvus = False
                wrote_vdb = False
                if vector_store.available():
                    mrows = [
                        {"id": rows[i].id, "source_id": source_id,
                         "text": chunks[i], "vector": list(vectors[i])}
                        for i in range(len(rows))
                    ]
                    wrote_vdb = vector_store.upsert(mrows)
                if wrote_vdb:
                    # 向量 + 文本入 LanceDB（它建 BM25 与向量索引）；SQLite 仅记录模型，不冗余存 bytes
                    for row in rows:
                        row.embedding_model = model
                    status = "ready"
                else:
                    # LanceDB 不可用/写失败 → 回落把向量存进 SQLite（纯 Python 暴力检索路径）
                    for i, row in enumerate(rows):
                        row.embedding = _vec_to_bytes(vectors[i])
                        row.embedding_model = model
                    status = "ready"

            source.chunk_count = len(chunks)
            source.index_status = status
            db.commit()
            return {"status": status, "chunks": len(chunks)}
        except Exception as e:
            source.index_status = "failed"
            source.index_error = str(e)
            db.commit()
            raise
    finally:
        db.close()


# ---------------------------------------------------------------- 混合检索 + 重排

# 召回候选池大小（每路各取这么多再融合）；RRF 常数 k=60 是论文/业界默认
CANDIDATE_K = 20
RRF_K = 60


def _top_indices(scores: list[float], pool: int) -> list[int]:
    """按分数降序取前 pool 个的下标（分数 > 0 才要），返回排名顺序的下标列表。"""
    order = sorted(range(len(scores)), key=lambda i: scores[i], reverse=True)
    out = []
    for i in order:
        if scores[i] <= 0:
            break
        out.append(i)
        if len(out) >= pool:
            break
    return out


def _rrf_fuse(rank_lists: list[list[int]], weights: list[float] = None, k: int = RRF_K) -> dict:
    """Reciprocal Rank Fusion：用「排名」而非「原始分数」融合多路召回。

    每个候选在某一路里排名 r（从 0 起）贡献 w / (k + r + 1)；累加得到融合分。
    优点：BM25 分与余弦相似度量纲完全不同、无法直接相加，而 RRF 只看名次，天然可比、稳健。
    返回 {chunk_index: 融合分}。
    """
    weights = weights or [1.0] * len(rank_lists)
    fused: dict[int, float] = {}
    for rl, w in zip(rank_lists, weights):
        for rank, idx in enumerate(rl):
            fused[idx] = fused.get(idx, 0.0) + w / (k + rank + 1)
    return fused


# ---------------------------------------------------------------- 检索

def retrieve(query: str, source_ids: list[str], top_k: int = 5,
             bm25_weight: float = 1.0, vector_weight: float = 1.0) -> list[dict]:
    """混合检索（hybrid retrieval + rerank）：

      路径 A（装了 LanceDB）：直接用 LanceDB 的混合检索 —— 全文 BM25 + 向量 + 内置 RRF 重排。
      路径 B（无 LanceDB）：纯 Python 回落 —— BM25 候选 + SQLite 暴力余弦候选 + 手写 RRF 融合。
    无 embedding 能力（如 Anthropic）时只走 BM25/全文；两路都没有则返回空。
    bm25_weight / vector_weight 仅作用于路径 B 的 RRF 加权。
    """
    if not source_ids:
        return []
    db = SessionLocal()
    try:
        chunks = (
            db.query(KnowledgeChunk, KnowledgeSource)
            .join(KnowledgeSource, KnowledgeChunk.source_id == KnowledgeSource.id)
            .filter(KnowledgeChunk.source_id.in_(source_ids))
            .all()
        )
        if not chunks:
            return []

        n = len(chunks)
        texts = [c.text for c, _ in chunks]
        pool = max(top_k * 4, CANDIDATE_K)
        cmap = {c.id: (c, s) for c, s in chunks}  # chunk_id → (chunk, source)

        def _fmt(score, c, s, backend):
            return {
                "source_id": c.source_id,
                "source_name": s.name,
                "category": s.category,
                "prompt": s.prompt,
                "text": c.text,
                "score": round(float(score), 6),
                "match": backend,  # lancedb-hybrid / lancedb-fts / py-hybrid / py-bm25（调试用）
            }

        # query embedding（无 embedding 能力的厂商会返回 None）
        try:
            qres = embed_texts([query])
        except Exception as e:
            print(f"[RAG] query embedding failed: {e}")
            qres = None
        qvec = qres[0][0] if qres else None

        # ========== 路径 A：专门向量库 LanceDB（BM25 + 向量 + 内置 RRF）==========
        if vector_store.available():
            try:
                if qvec is not None:
                    hits = vector_store.hybrid_search(query, qvec, source_ids, top_k)
                    backend = "lancedb-hybrid"
                else:
                    hits = vector_store.fts_search(query, source_ids, top_k)
                    backend = "lancedb-fts"
                if hits:
                    return [_fmt(sc, *cmap[cid], backend) for cid, sc in hits if cid in cmap]
            except Exception as e:
                print(f"[RAG] lancedb retrieve failed, fallback pure-python: {e}")

        # ========== 路径 B：纯 Python 回落（BM25 候选 + 暴力余弦候选 + RRF）==========
        bm25_scores = _bm25_rank(query, texts)
        bm25_rank = _top_indices(bm25_scores, pool)

        vec_rank: list[int] = []
        if qvec is not None and all(c.embedding is not None for c, _ in chunks):
            cos_scores = [_cosine(qvec, _bytes_to_vec(c.embedding)) for c, _ in chunks]
            vec_rank = _top_indices(cos_scores, pool)

        if vec_rank:
            fused = _rrf_fuse([bm25_rank, vec_rank], weights=[bm25_weight, vector_weight])
            ranked = sorted(fused.items(), key=lambda kv: kv[1], reverse=True)[:top_k]
            return [_fmt(fs, *chunks[idx], "py-hybrid") for idx, fs in ranked]

        # 只有 BM25 一路
        return [_fmt(bm25_scores[idx], *chunks[idx], "py-bm25") for idx in bm25_rank[:top_k]]
    finally:
        db.close()


def get_tail_text(source_id: str, max_chars: int = 3000) -> str:
    """续写专用：取文档末尾内容（按 chunk_index 倒序拼接）。"""
    db = SessionLocal()
    try:
        rows = (
            db.query(KnowledgeChunk)
            .filter(KnowledgeChunk.source_id == source_id)
            .order_by(KnowledgeChunk.chunk_index.desc())
            .all()
        )
        parts, total = [], 0
        for row in rows:
            parts.append(row.text)
            total += len(row.text)
            if total >= max_chars:
                break
        return "\n".join(reversed(parts))[-max_chars:]
    finally:
        db.close()


# ---------------------------------------------------------------- 写作卡上下文组装

SECTION_LABELS = {
    "writing_guide": "写作指导",
    "style": "风格参考",
    "reference": "资料库",
    "continuation": "续写原文",
}


def build_card_context(card, query: str, per_category_k: int = 4, max_chars: int = 9000) -> dict:
    """根据写作卡 + 当前查询（书籍需求或章节概要）组装各分类上下文。

    返回 {"context_block": str, "extra_requirements": str, "has_continuation": bool}
    context_block 可直接拼入 system prompt。
    """
    sections = []
    has_continuation = bool(card.continuation_ids)

    def render(category: str, ids: list[str], chunks: list[dict], extra_head: str = ""):
        if not chunks and not extra_head:
            return
        db = SessionLocal()
        try:
            prompts = [
                s.prompt
                for s in db.query(KnowledgeSource)
                .filter(KnowledgeSource.id.in_(ids))
                .all()
                if s.prompt
            ]
        finally:
            db.close()
        body = []
        if prompts:
            body.append("使用要求：" + "；".join(prompts))
        if extra_head:
            body.append(extra_head)
        for ch in chunks:
            body.append(f"〔{ch['source_name']}〕{ch['text']}")
        text = "\n".join(body)
        if len(text) > max_chars:
            text = text[:max_chars] + "…"
        sections.append(f"## {SECTION_LABELS[category]}\n{text}")

    if card.writing_guide_ids:
        render("writing_guide", card.writing_guide_ids, retrieve(query, card.writing_guide_ids, per_category_k))
    if card.style_ids:
        render("style", card.style_ids, retrieve(query, card.style_ids, per_category_k))
    if card.reference_ids:
        render("reference", card.reference_ids, retrieve(query, card.reference_ids, per_category_k))
    if card.continuation_ids:
        tails = [get_tail_text(sid) for sid in card.continuation_ids]
        tail_text = "\n---\n".join(t for t in tails if t)
        extra_head = (
            "以下是待续写文章的结尾部分（必须从此处自然衔接，延续其情节、人物、设定与文风）：\n"
            + tail_text
        )
        render(
            "continuation",
            card.continuation_ids,
            retrieve(query, card.continuation_ids, per_category_k),
            extra_head=extra_head,
        )

    context_block = "\n\n".join(sections)
    return {
        "context_block": context_block,
        "extra_requirements": (card.extra_requirements or "").strip(),
        "has_continuation": has_continuation,
    }


def new_id() -> str:
    return uuid.uuid4().hex[:16]
