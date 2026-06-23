import uuid
import json
import re
import io
import os
import html
import tempfile
from pathlib import Path
from urllib.parse import quote
from datetime import datetime
from sqlalchemy.orm import Session
from models import Book
from config import BOOKS_DIR, ensure_books_dir
from services.cover_gen import try_generate_cover

WINDOWS_RESERVED = {
    'con', 'prn', 'aux', 'nul',
    'com1', 'com2', 'com3', 'com4', 'com5', 'com6', 'com7', 'com8', 'com9',
    'lpt1', 'lpt2', 'lpt3', 'lpt4', 'lpt5', 'lpt6', 'lpt7', 'lpt8', 'lpt9',
}

def sanitize_filename(name: str) -> str:
    """清理文件名，移除非法字符，兼容多语言"""
    name = re.sub(r'[\x00-\x1f\x7f]', '', name)
    name = re.sub(r'[<>:"/\\|?*]', '', name)
    name = name.strip(' .').strip()[:100]
    if name.lower() in WINDOWS_RESERVED:
        name = f"_{name}"
    return name if name else "untitled"

def safe_filename_header(title: str, ext: str) -> str:
    """生成 Content-Disposition header，同时提供 ASCII fallback 和 UTF-8 编码"""
    safe = sanitize_filename(title)
    ascii_safe = re.sub(r'[^\x20-\x7e]', '', safe).strip(' .') or "book"
    encoded = quote(safe)
    return f'attachment; filename="{ascii_safe}.{ext}"; filename*=UTF-8\'\'{encoded}.{ext}'

def get_book_filepath(book_id: str, title: str) -> Path:
    """获取书籍文件保存路径"""
    ensure_books_dir()
    safe_title = sanitize_filename(title)
    filename = f"{book_id}_{safe_title}.md"
    return Path(BOOKS_DIR) / filename

def _chapter_prefix(index: int, language: str = "zh-CN") -> str:
    lang = (language or "").lower()
    if lang.startswith("zh"):
        return f"第{index}章："
    elif lang.startswith("ja"):
        return f"第{index}章："
    elif lang.startswith("ko"):
        return f"제{index}장: "
    elif lang == "en":
        return f"Chapter {index}: "
    elif lang.startswith("es"):
        return f"Capítulo {index}: "
    elif lang.startswith("fr"):
        return f"Chapitre {index}: "
    elif lang.startswith("de"):
        return f"Kapitel {index}: "
    elif lang.startswith("pt"):
        return f"Capítulo {index}: "
    elif lang.startswith("ru"):
        return f"Глава {index}: "
    elif lang.startswith("ar"):
        return f"الفصل {index}: "
    else:
        return f"Chapter {index}: "

_CJK_CHAR_RE = re.compile(r'[一-鿿぀-ヿ가-힯]')
_LATIN_WORD_RE = re.compile(r'[a-zA-Zà-ÿÀ-Ÿа-яА-Я0-9]+')


def count_words(text: str) -> int:
    """统计字数：CJK 按字符计，其他按单词计。"""
    plain = re.sub(r'[#*`>\[\]()_~-]', '', text)
    cjk = len(_CJK_CHAR_RE.findall(plain))
    latin = len(_LATIN_WORD_RE.findall(_CJK_CHAR_RE.sub(' ', plain)))
    return cjk + latin


def _do_save_book(
    db: Session,
    outline: dict,
    chapters: list,
    user_id: str,
    language: str = "zh-CN",
    tags: list = None,
) -> str:
    book_id = str(uuid.uuid4())  # 完整 UUID

    # 生成 Markdown 内容
    md_content = f"# {outline['title']}\n\n"
    md_content += f"{outline['description']}\n\n---\n\n"

    for i, (chapter_data, content) in enumerate(zip(outline['chapters'], chapters)):
        md_content += f"## {_chapter_prefix(i+1, language)}{chapter_data['title']}\n\n{content}\n\n"

    # 保存到本地文件
    filepath = get_book_filepath(book_id, outline['title'])
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(md_content)

    # 标签：调用方传入（用户标签 + AI 标签合并），兜底取大纲里的 AI 标签
    final_tags = tags if tags else outline.get("tags") or []
    if not isinstance(final_tags, list):
        final_tags = []
    # 统一规范：去掉用户/AI 可能带的 # 前缀，存纯文字（展示层负责加 #）
    final_tags = [str(t).strip().lstrip("#").strip() for t in final_tags]
    final_tags = [t for t in final_tags if t][:10]

    # 保存到数据库
    book = Book(
        id=book_id,
        title=outline['title'],
        description=outline['description'],
        outline=outline,
        file_path=str(filepath),
        author_id=user_id,
        source="local",
        language=language,
        tags=final_tags,
        word_count=count_words(md_content),
    )
    db.add(book)
    db.commit()

    return book_id


def save_book(db: Session, outline: dict, chapters: list, user_id: str, language: str = "zh-CN", tags: list = None) -> str:
    """保存书籍到本地文件和数据库"""
    return _do_save_book(db, outline, chapters, user_id, language, tags)

def save_book_sync(db: Session, outline: dict, chapters: list, user_id: str, language: str = "zh-CN", tags: list = None) -> str:
    """同步版本的书籍保存，用于后台线程"""
    return _do_save_book(db, outline, chapters, user_id, language, tags)

def save_p2p_book(db: Session, book_data: dict, peer_id: str) -> str:
    """保存 P2P 接收的书籍"""
    book_id = book_data['id']
    
    # 生成 Markdown 内容
    md_content = f"# {book_data['title']}\n\n"
    md_content += f"{book_data['description']}\n\n---\n\n"
    
    for ch in book_data.get('chapters', []):
        md_content += f"## {ch['title']}\n\n{ch['content']}\n\n"
    
    # 保存到本地文件
    filepath = get_book_filepath(book_id, book_data['title'])
    with open(filepath, "w", encoding="utf-8") as f:
        f.write(md_content)
    
    # 保存到数据库
    book = Book(
        id=book_id,
        title=book_data['title'],
        description=book_data['description'],
        outline=book_data.get('outline', {}),
        file_path=str(filepath),
        source="p2p",
        peer_origin=peer_id,
        language=book_data.get('language')
    )
    db.add(book)
    db.commit()
    
    return book_id

def get_book(db: Session, book_id: str):
    return db.query(Book).filter(Book.id == book_id).first()

def get_all_books(db: Session):
    return db.query(Book).order_by(Book.created_at.desc()).all()

def delete_book(db: Session, book_id: str) -> bool:
    book = db.query(Book).filter(Book.id == book_id).first()
    if book:
        # 删除本地文件
        filepath = Path(book.file_path)
        if filepath.exists():
            filepath.unlink()
        
        db.delete(book)
        db.commit()
        return True
    return False

def get_book_content(db: Session, book_id: str) -> str:
    """读取书籍文件内容"""
    book = get_book(db, book_id)
    if not book:
        return None
    
    filepath = Path(book.file_path)
    if not filepath.exists():
        return None
    
    with open(filepath, "r", encoding="utf-8") as f:
        return f.read()

def get_book_chapters(db: Session, book_id: str) -> list:
    """从 Markdown 内容解析章节"""
    content = get_book_content(db, book_id)
    if not content:
        return []
    
    chapters = []
    # 按 ## 分割章节
    parts = re.split(r'\n## ', content)
    
    for i, part in enumerate(parts):
        if i == 0:
            continue  # 跳过标题部分
        
        lines = part.strip().split('\n', 1)
        title = lines[0].strip()
        chapter_content = lines[1].strip() if len(lines) > 1 else ""
        
        # 移除章节标题中的"第X章："前缀
        title = re.sub(r'^(第\d+章[：:]?\s*|Chapter\s+\d+[\s:]*|Capítulo\s+\d+[\s:]*|Chapitre\s+\d+[\s:]*|Kapitel\s+\d+[\s:]*|Глава\s+\d+[\s:]*)', '', title).strip()
        
        chapters.append({
            'index': i - 1,
            'title': title,
            'content': chapter_content
        })
    
    return chapters

def export_to_markdown(db: Session, book_id: str) -> str:
    """导出为 Markdown"""
    return get_book_content(db, book_id)

def _parse_chapters(md_content: str):
    """从 Markdown 内容解析出标题、描述和章节列表"""
    title = ""
    description = ""
    chapters = []

    chapter_parts = re.split(r'\n## ', md_content)

    for i, part in enumerate(chapter_parts):
        if i == 0:
            header_lines = part.strip().split('\n')
            for line in header_lines:
                if line.startswith('# '):
                    title = line[2:].strip()
                elif line.strip() and line.strip() != '---':
                    if not description:
                        description = line.strip()
                    else:
                        description += '\n' + line.strip()
        else:
            lines_text = part.strip().split('\n', 1)
            ch_title = lines_text[0].strip()
            ch_content = lines_text[1].strip() if len(lines_text) > 1 else ""
            ch_title = re.sub(r'^(第\d+章[：:]?\s*|Chapter\s+\d+[\s:]*|Capítulo\s+\d+[\s:]*|Chapitre\s+\d+[\s:]*|Kapitel\s+\d+[\s:]*|Глава\s+\d+[\s:]*)', '', ch_title).strip()
            chapters.append({"title": ch_title, "content": ch_content})

    return title, description, chapters


def export_to_txt(md_content: str) -> bytes:
    """导出为纯文本，去除 Markdown 格式标记"""
    text = md_content
    text = re.sub(r'```[\s\S]*?```', lambda m: m.group(0).replace('```', '').strip('\n'), text)
    text = re.sub(r'!\[([^\]]*)\]\([^)]+\)', r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'\*{1,2}([^*]+)\*{1,2}', r'\1', text)
    text = re.sub(r'_{1,2}([^_]+)_{1,2}', r'\1', text)
    text = re.sub(r'~~([^~]+)~~', r'\1', text)
    text = re.sub(r'^\s*[-*+]\s+', '  - ', text, flags=re.MULTILINE)
    text = re.sub(r'^\s*\d+\.\s+', '  ', text, flags=re.MULTILINE)
    text = re.sub(r'^---+\s*$', '', text, flags=re.MULTILINE)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip().encode('utf-8')


def export_to_epub(md_content: str, title: str = "Book", language: str = "zh-CN", book_id: str = "") -> bytes:
    """导出为 EPUB"""
    from ebooklib import epub

    book_title, description, chapters = _parse_chapters(md_content)
    if title and title != "Book":
        book_title = title

    epub_lang = (language or "zh").split("-")[0]

    book = epub.EpubBook()
    book.set_identifier(str(uuid.uuid4()))
    book.set_title(book_title)
    book.set_language(epub_lang)
    book.add_author('AI eBook Generator')
    if description:
        book.add_metadata('DC', 'description', html.escape(description))

    # 封面：生成算法封面并设为 EPUB 封面（阅读器会显示在首页/书架）
    cover_bytes = try_generate_cover(book_title, book_id)
    has_cover = False
    if cover_bytes:
        try:
            book.set_cover("cover.png", cover_bytes)
            has_cover = True
        except Exception as e:
            print(f"[EPUB] set_cover failed: {e}")

    # 封面页放在最前，其次是导航
    spine_items = (['cover', 'nav'] if has_cover else ['nav'])
    toc_items = []

    style = '''
    body { font-family: serif; line-height: 1.8; margin: 1em; }
    h1 { font-size: 1.6em; text-align: center; margin: 1.5em 0 0.8em; }
    h2 { font-size: 1.3em; margin: 1.2em 0 0.6em; border-bottom: 1px solid #ddd; padding-bottom: 0.3em; }
    h3 { font-size: 1.1em; margin: 1em 0 0.4em; }
    p { margin: 0.5em 0; text-indent: 2em; }
    pre { background: #f5f5f5; padding: 1em; overflow-x: auto; font-size: 0.85em; }
    code { background: #f5f5f5; padding: 0.1em 0.3em; font-size: 0.9em; }
    blockquote { border-left: 3px solid #ccc; margin: 0.5em 0; padding: 0.5em 1em; color: #555; }
    '''

    nav_css = epub.EpubItem(uid='style_nav', file_name='style/nav.css', media_type='text/css', content=style)
    book.add_item(nav_css)

    intro_html = f'''<html><body>
    <h1>{html.escape(book_title)}</h1>
    {'<p>' + html.escape(description) + '</p>' if description else ''}
    </body></html>'''
    intro = epub.EpubHtml(title=book_title, file_name='intro.xhtml', lang=epub_lang)
    intro.content = intro_html
    intro.add_item(nav_css)
    book.add_item(intro)
    spine_items.append(intro)

    for i, ch in enumerate(chapters):
        ch_html_content = _markdown_to_simple_html(ch['content'])
        chapter = epub.EpubHtml(
            title=html.escape(ch['title']),
            file_name=f'chapter_{i+1}.xhtml',
            lang=epub_lang
        )
        chapter.content = f'''<html><head><link rel="stylesheet" href="style/nav.css"/></head>
        <body><h2>{html.escape(ch['title'])}</h2>{ch_html_content}</body></html>'''
        chapter.add_item(nav_css)
        book.add_item(chapter)
        spine_items.append(chapter)
        toc_items.append(chapter)

    book.toc = toc_items
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    book.spine = spine_items

    buffer = io.BytesIO()
    epub.write_epub(buffer, book)
    buffer.seek(0)
    return buffer.read()


def export_to_docx(md_content: str, title: str = "Book", language: str = "zh-CN", book_id: str = "") -> bytes:
    """导出为 Word (.docx)"""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    book_title, description, chapters = _parse_chapters(md_content)
    if title and title != "Book":
        book_title = title

    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    heading_style = doc.styles['Heading 1']
    heading_style.font.size = Pt(22)
    heading_style.font.bold = True

    heading2_style = doc.styles['Heading 2']
    heading2_style.font.size = Pt(16)
    heading2_style.font.bold = True

    # 封面页：整页居中的封面图，单独成页
    cover_bytes = try_generate_cover(book_title, book_id)
    if cover_bytes:
        try:
            from docx.enum.text import WD_BREAK
            cover_para = doc.add_paragraph()
            cover_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cover_para.add_run().add_picture(io.BytesIO(cover_bytes), width=Inches(5.0))
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        except Exception as e:
            print(f"[DOCX] add cover failed: {e}")

    doc.add_heading(book_title, level=1)

    if description:
        desc_para = doc.add_paragraph(description)
        desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        desc_para.style.font.size = Pt(11)

    doc.add_paragraph('')

    for ch in chapters:
        doc.add_heading(ch['title'], level=2)
        for block in _parse_content_blocks(ch['content']):
            kind = block[0]
            if kind == "heading":
                _, level, text = block
                if level <= 3:
                    doc.add_heading(text, level=min(level, 2))
                else:
                    para = doc.add_paragraph()
                    para.add_run(text).bold = True
            elif kind == "list":
                for item in block[1]:
                    doc.add_paragraph(_strip_inline_md(item), style='List Bullet')
            elif kind == "code":
                code_para = doc.add_paragraph(block[1])
                code_para.style.font.name = 'Courier New'
                code_para.style.font.size = Pt(9)
            elif kind == "table":
                headers, rows = block[1], block[2]
                if headers:
                    tbl = doc.add_table(rows=1, cols=len(headers))
                    tbl.style = 'Light Grid Accent 1'
                    for c, h in enumerate(headers):
                        tbl.rows[0].cells[c].text = _strip_inline_md(h)
                    for row in rows:
                        cells = (row + [''] * len(headers))[:len(headers)]
                        wr = tbl.add_row().cells
                        for c, val in enumerate(cells):
                            wr[c].text = _strip_inline_md(val)
                    doc.add_paragraph('')
            elif kind == "svg":
                png = _rasterize_svg_to_png(block[1])
                if png:
                    try:
                        pic = doc.add_paragraph()
                        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        pic.add_run().add_picture(io.BytesIO(png), width=Inches(5.5))
                    except Exception:
                        png = None
                if not png and block[2]:
                    cap = doc.add_paragraph(f"〔图：{block[2]}〕")
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if block[2]:
                    cap = doc.add_paragraph(block[2])
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap.runs and setattr(cap.runs[0].font, 'size', Pt(9))
            else:  # para
                doc.add_paragraph(_strip_inline_md(block[1]))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def export_to_pdf(md_content: str, title: str = "Book", language: str = "zh-CN", book_id: str = "") -> bytes:
    """导出为 PDF（使用 reportlab，纯 Python 无系统依赖）"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Image
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.fonts import addMapping

    book_title, description, chapters = _parse_chapters(md_content)
    if title and title != "Book":
        book_title = title

    font_name = _register_pdf_fonts()

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'BookTitle', parent=styles['Title'],
        fontSize=22, leading=28, alignment=TA_CENTER,
        spaceAfter=12, fontName=font_name,
    )
    desc_style = ParagraphStyle(
        'BookDesc', parent=styles['Normal'],
        fontSize=12, leading=16, alignment=TA_CENTER,
        textColor='#666666', spaceAfter=20, fontName=font_name,
    )
    h2_style = ParagraphStyle(
        'BookH2', parent=styles['Heading2'],
        fontSize=16, leading=20, spaceAfter=8, spaceBefore=20,
        fontName=font_name,
    )
    body_style = ParagraphStyle(
        'BookBody', parent=styles['Normal'],
        fontSize=11, leading=18, spaceAfter=6,
        firstLineIndent=22, fontName=font_name,
    )
    code_style = ParagraphStyle(
        'BookCode', parent=styles['Code'],
        fontSize=9, leading=13, fontName='Courier',
        backColor='#F5F5F5', leftIndent=10, rightIndent=10,
        spaceBefore=6, spaceAfter=6,
    )

    story = []

    # 封面页：整页封面图，独立成页
    cover_bytes = try_generate_cover(book_title, book_id)
    if cover_bytes:
        try:
            avail_w = doc.width
            avail_h = doc.height
            img_w = avail_w
            img_h = img_w * 800 / 600
            if img_h > avail_h:
                img_h = avail_h
                img_w = img_h * 600 / 800
            cover_img = Image(io.BytesIO(cover_bytes), width=img_w, height=img_h)
            cover_img.hAlign = 'CENTER'
            story.append(cover_img)
            story.append(PageBreak())
        except Exception as e:
            print(f"[PDF] add cover failed: {e}")

    story.append(Paragraph(_pdf_escape(book_title), title_style))
    story.append(Spacer(1, 6))
    if description:
        story.append(Paragraph(_pdf_escape(description), desc_style))
        story.append(Spacer(1, 12))

    for ch in chapters:
        story.append(PageBreak())
        story.append(Paragraph(_pdf_escape(ch['title']), h2_style))
        story.append(Spacer(1, 6))

        from reportlab.platypus import Table, TableStyle
        from reportlab.lib import colors
        bullet_style = ParagraphStyle('Bullet', parent=body_style, firstLineIndent=0, leftIndent=20, bulletIndent=10)
        cap_style = ParagraphStyle('Cap', parent=body_style, firstLineIndent=0, alignment=TA_CENTER, fontSize=9, textColor='#666666')

        for block in _parse_content_blocks(ch['content']):
            kind = block[0]
            if kind == "heading":
                _, level, text = block
                story.append(Paragraph(_pdf_escape(text), h2_style if level <= 2 else body_style))
            elif kind == "list":
                for item in block[1]:
                    story.append(Paragraph(_pdf_escape(_strip_inline_md(item)), bullet_style, bulletText='•'))
            elif kind == "code":
                story.append(Paragraph(_pdf_escape(block[1]), code_style))
            elif kind == "table":
                headers, rows = block[1], block[2]
                if headers:
                    data = [[Paragraph(_pdf_escape(_strip_inline_md(h)), body_style) for h in headers]]
                    for row in rows:
                        cells = (row + [''] * len(headers))[:len(headers)]
                        data.append([Paragraph(_pdf_escape(_strip_inline_md(c)), body_style) for c in cells])
                    tbl = Table(data, hAlign='LEFT')
                    tbl.setStyle(TableStyle([
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#eef2ff')),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                        ('LEFTPADDING', (0, 0), (-1, -1), 5),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 5),
                        ('TOPPADDING', (0, 0), (-1, -1), 3),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
                    ]))
                    story.append(Spacer(1, 6))
                    story.append(tbl)
                    story.append(Spacer(1, 8))
            elif kind == "svg":
                drawn = False
                try:
                    from io import BytesIO as _BIO
                    from svglib.svglib import svg2rlg
                    drawing = svg2rlg(_BIO(block[1].encode("utf-8")))
                    if drawing is not None:
                        story.append(Spacer(1, 6))
                        story.append(drawing)
                        drawn = True
                except Exception:
                    drawn = False
                if not drawn:
                    png = _rasterize_svg_to_png(block[1])
                    if png:
                        try:
                            story.append(Image(io.BytesIO(png), width=14 * cm, height=14 * cm * 380 / 640))
                            drawn = True
                        except Exception:
                            drawn = False
                if block[2]:
                    story.append(Paragraph(_pdf_escape(f"图：{block[2]}"), cap_style))
                story.append(Spacer(1, 8))
            else:  # para
                story.append(Paragraph(_pdf_escape(_strip_inline_md(block[1])), body_style))

    doc.build(story)
    buffer.seek(0)
    return buffer.read()


def _register_pdf_fonts():
    """注册 PDF 字体，优先使用系统中文字体"""
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    candidates = []
    if os.name == 'nt':
        windir = os.environ.get('WINDIR', r'C:\Windows')
        fonts_dir = os.path.join(windir, 'Fonts')
        candidates = [
            (os.path.join(fonts_dir, 'msyh.ttc'), 'msyh', 0),
            (os.path.join(fonts_dir, 'msyhbd.ttc'), 'msyhbd', 0),
            (os.path.join(fonts_dir, 'msyhl.ttc'), 'msyhl', 0),
            (os.path.join(fonts_dir, 'simhei.ttf'), 'simhei', -1),
            (os.path.join(fonts_dir, 'simsun.ttc'), 'simsun', 0),
        ]
    else:
        candidates = [
            ('/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc', 'notocjk', 0),
            ('/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc', 'notocjk', 0),
            ('/usr/share/fonts/noto-cjk/NotoSansCJK-Regular.ttc', 'notocjk', 0),
            ('/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc', 'wqy', 0),
            ('/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf', 'droid', -1),
        ]

    for path, name, idx in candidates:
        try:
            if os.path.exists(path):
                pdfmetrics.registerFont(TTFont(name, path, subfontIndex=idx))
                return name
        except Exception:
            continue

    return 'Helvetica'


def _pdf_escape(text: str) -> str:
    """转义 reportlab Paragraph 中的特殊字符"""
    if not text:
        return ''
    text = html.escape(text, quote=False)
    text = text.replace('\n', '<br/>')
    return text


def _gfm_table_to_html(block: str) -> str:
    """把一段 GFM 表格文本转成 <table>。非表格返回 None。"""
    lines = [l for l in block.strip().split('\n') if l.strip()]
    if len(lines) < 2 or not re.match(r'^\s*\|?\s*:?-{2,}', lines[1].replace(' ', '')):
        # 第二行须是分隔行（---|---）
        if len(lines) < 2 or '---' not in lines[1]:
            return None

    def cells(line):
        line = line.strip().strip('|')
        return [c.strip() for c in line.split('|')]

    headers = cells(lines[0])
    body = [cells(l) for l in lines[2:]]
    out = ['<table border="1" cellspacing="0" cellpadding="4" style="border-collapse:collapse;margin:1em 0;">']
    out.append('<thead><tr>' + ''.join(f'<th>{html.escape(h)}</th>' for h in headers) + '</tr></thead>')
    out.append('<tbody>')
    for row in body:
        row = (row + [''] * len(headers))[:len(headers)]
        out.append('<tr>' + ''.join(f'<td>{html.escape(c)}</td>' for c in row) + '</tr>')
    out.append('</tbody></table>')
    return ''.join(out)


def _rasterize_svg_to_png(svg: str):
    """SVG → PNG bytes，供 DOCX/PDF 嵌入。无可用渲染器时返回 None（优雅降级为说明文字）。"""
    if not svg:
        return None
    try:
        import cairosvg
        return cairosvg.svg2png(bytestring=svg.encode("utf-8"), output_width=640)
    except Exception:
        pass
    try:
        from io import BytesIO
        from svglib.svglib import svg2rlg
        from reportlab.graphics import renderPM
        drawing = svg2rlg(BytesIO(svg.encode("utf-8")))
        if drawing is None:
            return None
        buf = BytesIO()
        renderPM.drawToFile(drawing, buf, fmt="PNG")
        return buf.getvalue()
    except Exception:
        return None


def _parse_content_blocks(text: str):
    """把章节 Markdown 解析为有序块列表，供 DOCX/PDF 渲染。

    返回元素：
      ("heading", level, text) / ("svg", svg_str, caption) / ("table", headers, rows)
      ("code", code_text) / ("list", [items]) / ("para", text)
    """
    lines = text.split("\n")
    blocks = []
    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        # SVG 围栏
        if stripped.startswith("```svg"):
            j = i + 1
            buf = []
            while j < n and lines[j].strip() != "```":
                buf.append(lines[j])
                j += 1
            svg = "\n".join(buf).strip()
            caption = ""
            k = j + 1
            while k < n and not lines[k].strip():
                k += 1
            if k < n:
                cm = re.match(r'^\*图[:：]\s*(.+?)\*$', lines[k].strip())
                if cm:
                    caption = cm.group(1)
                    k += 1
            blocks.append(("svg", svg, caption))
            i = k
            continue

        # 其它代码围栏
        if stripped.startswith("```"):
            j = i + 1
            buf = []
            while j < n and lines[j].strip() != "```":
                buf.append(lines[j])
                j += 1
            blocks.append(("code", "\n".join(buf)))
            i = j + 1
            continue

        # GFM 表格：当前行含 |，下一行是分隔行
        if "|" in stripped and i + 1 < n and re.match(r'^\s*\|?[\s:\-|]+\|?\s*$', lines[i + 1]) and "-" in lines[i + 1]:
            def cells(l):
                return [c.strip() for c in l.strip().strip("|").split("|")]
            headers = cells(line)
            rows = []
            j = i + 2
            while j < n and "|" in lines[j] and lines[j].strip():
                rows.append(cells(lines[j]))
                j += 1
            blocks.append(("table", headers, rows))
            i = j
            continue

        # 标题
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            blocks.append(("heading", level, stripped.lstrip("#").strip()))
            i += 1
            continue

        # 列表（连续项）
        if stripped.startswith(("- ", "* ")):
            items = []
            while i < n and lines[i].strip().startswith(("- ", "* ")):
                items.append(lines[i].strip()[2:])
                i += 1
            blocks.append(("list", items))
            continue

        # 普通段落
        blocks.append(("para", stripped))
        i += 1
    return blocks


def _strip_inline_md(s: str) -> str:
    s = re.sub(r'`⚓[^`]*`', '', s)          # 移除 stub 锚点 token
    s = re.sub(r'\*{1,2}([^*]+)\*{1,2}', r'\1', s)
    s = re.sub(r'_{1,2}([^_]+)_{1,2}', r'\1', s)
    s = re.sub(r'~~([^~]+)~~', r'\1', s)
    s = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', s)
    s = re.sub(r'`([^`]+)`', r'\1', s)
    return s


def _markdown_to_simple_html(md_text: str) -> str:
    """将 Markdown 章节内容转为简单 HTML（用于 EPUB 章节体）"""
    import html as html_mod

    html = md_text

    # 0) stub 锚点 token ``⚓ID`` → 隐藏 <a id>；跨章跳转链接降级为纯文字（EPUB 章节分文件，跨文件锚点不可靠）
    html = re.sub(r'`⚓([^`]+)`', lambda m: f'<a id="{html_mod.escape(m.group(1))}"></a>', html)
    html = re.sub(r'（?\[([^\]]+)\]\(#[^)]+\)）?', r'\1', html)

    # 1) 写作工具 SVG：```svg ... ``` 直接内联（EPUB3/XHTML 支持 SVG）
    html = re.sub(r'```svg\s*\n([\s\S]*?)```', lambda m: f'<div style="text-align:center;margin:1em 0;">{m.group(1).strip()}</div>', html)

    # 2) GFM 表格 → <table>
    def _table_repl(m):
        res = _gfm_table_to_html(m.group(0))
        return res if res else m.group(0)
    html = re.sub(r'(?:^\|.*\|[ \t]*\n)(?:^\|?[ \t:\-|]+\|?[ \t]*\n)(?:^\|.*\|[ \t]*\n?)+', _table_repl, html, flags=re.MULTILINE)

    html = re.sub(r'```(\w*)\n([\s\S]*?)```', r'<pre><code>\2</code></pre>', html)
    html = re.sub(r'!\[([^\]]*)\]\([^)]+\)', r'', html)
    html = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', html)
    html = re.sub(r'^####\s+(.+)$', r'<h4>\1</h4>', html, flags=re.MULTILINE)
    html = re.sub(r'^###\s+(.+)$', r'<h3>\1</h3>', html, flags=re.MULTILINE)
    html = re.sub(r'^##\s+(.+)$', r'<h2>\1</h2>', html, flags=re.MULTILINE)
    html = re.sub(r'^#\s+(.+)$', r'<h1>\1</h1>', html, flags=re.MULTILINE)
    html = re.sub(r'^\*\s+(.+)$', r'<li>\1</li>', html, flags=re.MULTILINE)
    html = re.sub(r'^-\s+(.+)$', r'<li>\1</li>', html, flags=re.MULTILINE)
    html = re.sub(r'^>\s+(.+)$', r'<blockquote>\1</blockquote>', html, flags=re.MULTILINE)
    html = re.sub(r'\*\*([^*]+)\*\*', r'<strong>\1</strong>', html)
    html = re.sub(r'\*([^*]+)\*', r'<em>\1</em>', html)
    html = re.sub(r'~~([^~]+)~~', r'<del>\1</del>', html)

    paragraphs = html.split('\n\n')
    result = []
    for p in paragraphs:
        p = p.strip()
        if not p:
            continue
        if p.startswith('<h') or p.startswith('<pre') or p.startswith('<li') or p.startswith('<blockquote') or p.startswith('<table') or p.startswith('<div'):
            result.append(p)
        else:
            lines = p.split('\n')
            processed = []
            for line in lines:
                line = line.strip()
                if line:
                    processed.append(line)
            if processed:
                result.append('<p>' + ' '.join(processed) + '</p>')

    return '\n'.join(result)


def update_books_dir(new_dir: str):
    """更新书籍保存目录"""
    os.environ['BOOKS_DIR'] = new_dir
    ensure_books_dir()
